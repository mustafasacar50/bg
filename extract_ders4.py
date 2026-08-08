import json
import hashlib
import docx

questions = []

def add_q(sentence, answer, hint, explanation=None, lesson="Ders 4", type_override=None):
    # Unique ID
    raw = sentence + answer + hint + str(type_override)
    qid = "q_auto_" + hashlib.md5(raw.encode('utf-8')).hexdigest()[:8]
    q_type = type_override if type_override else ("fill_in_the_blank" if "_____" in sentence else "translation")
    q = {
        "id": qid,
        "type": q_type,
        "sentence": sentence,
        "answer": answer,
        "hint": hint,
        "lesson": lesson
    }
    if explanation:
        q["explanation"] = explanation
    
    # Avoid exact duplicates
    if not any(x['sentence'] == q['sentence'] and x['answer'] == q['answer'] and x['type'] == q['type'] for x in questions):
        questions.append(q)

def parse_docx(file_path):
    doc = docx.Document(file_path)
    
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    i = 0
    while i < len(paragraphs):
        p = paragraphs[i]
        
        if p.startswith('BG:'):
            bg_text = p[3:].strip()
            
            # Look ahead for TR:
            if i + 1 < len(paragraphs) and paragraphs[i+1].startswith('TR:'):
                tr_text = paragraphs[i+1][3:].strip()
                # Remove parentheses if they exist (e.g. TR: (O 18 yaşındadır.))
                if tr_text.startswith('(') and tr_text.endswith(')'):
                    tr_text = tr_text[1:-1].strip()
                i += 1
                
                exp = None
                
                # Look ahead for "Not:" explanations
                if i + 1 < len(paragraphs) and paragraphs[i+1].startswith('Not:'):
                    exp = paragraphs[i+1][4:].strip()
                    i += 1
                
                # BG to TR
                add_q(
                    sentence=f"Çeviriniz (Bulgarcası: {bg_text})",
                    answer=tr_text,
                    hint="Türkçe karşılığını yazınız",
                    explanation=exp
                )
                # TR to BG
                add_q(
                    sentence=f"Çeviriniz (Türkçesi: {tr_text})",
                    answer=bg_text,
                    hint="Bulgarca karşılığını yazınız",
                    explanation=exp
                )
                if len(bg_text.split()) >= 2:
                    add_q(
                        sentence=f"Kelimeleri sıraya dizerek cümleyi kurunuz: {tr_text}",
                        answer=bg_text,
                        hint="Doğru sıralamayı bul",
                        explanation=exp,
                        type_override="scramble"
                    )
                
        i += 1

    # Now parse tables (Vocab)
    for table in doc.tables:
        for r_idx in range(len(table.rows)):
            cells = table.rows[r_idx].cells
            
            # Combine split cells for Ders 4 quirks
            raw_text = ' '.join([c.text.strip().replace('\n', ' ') for c in cells])
            if '|' in raw_text or len(cells) < 2:
                continue
                
            bg_col = cells[0].text.strip().replace('\n', ' ')
            tr_col = cells[1].text.strip().replace('\n', ' ')
            
            if len(cells) >= 3 and not tr_col:
                tr_col = cells[2].text.strip().replace('\n', ' ')
                
            # Special fix for split cells in Ders 4
            if len(cells) >= 2:
                if 'meyveli tatlı' in cells[1].text:
                    bg_col = cells[0].text.strip() + " " + cells[1].text.replace('meyveli tatlı', '').strip()
                    tr_col = 'meyveli tatlı'
                elif 'kuru meyveler' in cells[1].text:
                    bg_col = cells[0].text.strip() + " " + cells[1].text.replace('kuru meyveler', '').strip()
                    tr_col = 'kuru meyveler'
                elif 'kuru kayısılar' in cells[1].text:
                    bg_col = cells[0].text.strip() + " " + cells[1].text.replace('kuru kayısılar', '').strip()
                    tr_col = 'kuru kayısılar'

            if not bg_col or not tr_col or bg_col.lower() == 'bulgarca' or tr_col.lower() == 'türkçe':
                continue
                
            exp = None
            if len(bg_col) > 1 and len(tr_col) > 1:
                add_q(f"Çeviriniz (Türkçesi: {tr_col})", bg_col, "Bulgarca karşılığını yazınız", exp)
                add_q(f"Çeviriniz (Bulgarcası: {bg_col})", tr_col, "Türkçe karşılığını yazınız", exp)
                if len(bg_col.split()) >= 2:
                    add_q(f"Kelimeleri sıraya dizerek cümleyi kurunuz: {tr_col}", bg_col, "Doğru sıralamayı bul", exp, type_override="scramble")

parse_docx(r'D:\bulgarca_sınav_modulu\docs\balgoc___Bulgarca_A1_Ders_4_Turkce_Aciklamali_Not (1).docx')

with open('src/data/modules/balgoc___Bulgarca_A1_Ders_4.json', 'w', encoding='utf-8') as f:
    json.dump({
        "id": "balgoc___Bulgarca_A1_Ders_4",
        "title": "Ders 4: Ev, Eşyalar ve Yiyecekler",
        "description": "Hatice ve Emin okuma metinleri, var/yok kullanımı, sayılar, renkler, meyve/sebze ve ev eşyaları kelimeleri.",
        "questions": questions
    }, f, ensure_ascii=False, indent=2)

print(f"Generated {len(questions)} questions in src/data/modules/balgoc___Bulgarca_A1_Ders_4.json")
