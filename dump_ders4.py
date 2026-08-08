import docx

doc = docx.Document(r'D:\bulgarca_sınav_modulu\docs\balgoc___Bulgarca_A1_Ders_4_Turkce_Aciklamali_Not (1).docx')

with open('ders4_full.txt', 'w', encoding='utf-8') as f:
    for p in doc.paragraphs:
        if p.text.strip():
            f.write(p.text.strip() + '\n')
            
    f.write('\n--- TABLES ---\n')
    for table in doc.tables:
        for row in table.rows:
            f.write(' | '.join([cell.text.strip().replace('\n', ' ') for cell in row.cells]) + '\n')
        f.write('\n')
