import docx
doc = docx.Document(r'D:\bulgarca_sınav_modulu\docs\balgoc___Bulgarca_A1_Ders_1_2_Tam_Ceviri_Aciklamali.docx')
with open('D:/bulgarca_sınav_modulu/exam-app/ders1_2_temp.txt', 'w', encoding='utf-8') as f:
    for p in doc.paragraphs:
        if p.text.strip(): f.write(p.text.strip() + '\n')
    for table in doc.tables:
        for row in table.rows:
            f.write('\t'.join([cell.text.strip() for cell in row.cells]) + '\n')
