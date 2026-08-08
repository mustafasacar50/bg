import json
import hashlib
import re
import docx

questions = []

def add_q(sentence, answer, hint, explanation=None, lesson="Ders 3"):
    # Unique ID
    raw = sentence + answer + hint
    qid = "q_auto_" + hashlib.md5(raw.encode('utf-8')).hexdigest()[:8]
    q = {
        "id": qid,
        "type": "fill_in_the_blank" if "_____" in sentence else "translation",
        "sentence": sentence,
        "answer": answer,
        "hint": hint,
        "lesson": lesson
    }
    if explanation:
        q["explanation"] = explanation
    
    # Avoid exact duplicates
    if not any(x['sentence'] == q['sentence'] and x['answer'] == q['answer'] for x in questions):
        questions.append(q)

def get_phonetic_rule(tr_word, bg_word):
    rules = []
    tr_upper = tr_word.upper()
    bg_upper = bg_word.upper()
    
    if 'Ş' in tr_upper and 'Ш' in bg_upper: rules.append("Türkçedeki Ş sesi Bulgarcada genellikle Ш ile yazılır.")
    if 'J' in tr_upper and 'Ж' in bg_upper: rules.append("Türkçedeki J sesi Bulgarcada Ж ile yazılır.")
    if 'Ç' in tr_upper and 'Ч' in bg_upper: rules.append("Türkçedeki Ç sesi Bulgarcada genellikle Ч ile yazılır.")
    if 'Ö' in tr_upper and 'ЬО' in bg_upper: rules.append("Ö sesi bazı kelimelerde yumuşatma işaretli ьо ile gösterilir.")
    if 'Ü' in tr_upper and 'Ю' in bg_upper: rules.append("Ü sesi kimi kelimelerde Ю ile yazılır.")
    
    return " / ".join(rules) if rules else None

def is_cyrillic(text):
    return any('\u0400' <= c <= '\u04FF' for c in text)

def parse_docx(filepath):
    doc = docx.Document(filepath)
    
    # 1. Parse paragraphs for BG / TR / Not structure
    i = 0
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    while i < len(paragraphs):
        text = paragraphs[i]
        
        if text.startswith('BG:'):
            bg_text = text.replace('BG:', '').strip()
            tr_text = ""
            note_text = ""
            
            if i + 1 < len(paragraphs) and paragraphs[i+1].startswith('TR:'):
                tr_text = paragraphs[i+1].replace('TR:', '').strip()
                # Remove brackets if they exist
                if tr_text.startswith('(') and tr_text.endswith(')'):
                    tr_text = tr_text[1:-1]
                i += 1
                
                if i + 1 < len(paragraphs) and paragraphs[i+1].startswith('Not:'):
                    note_text = paragraphs[i+1].replace('Not:', '').strip()
                    i += 1
            
            if bg_text and tr_text:
                exp = note_text if note_text else get_phonetic_rule(tr_text, bg_text)
                
                # BG to TR
                add_q(
                    sentence=f"Çeviriniz (Bulgarcası: {bg_text})",
                    answer=tr_text,
                    hint="Türkçe karşılığını yazınız",
                    explanation=exp
                )
                # TR to BG
                add_q(
                    sentence=f"Çeviriniz (Türkçesi: {tr_text})",
                    answer=bg_text,
                    hint="Bulgarca karşılığını yazınız",
                    explanation=exp
                )
                
        i += 1

    # 2. Parse Tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            if not cells or len(cells) < 2:
                continue
                
            # Table Header checks
            if 'Bulgarca' in cells[0] or 'Türkçe' in cells[1] or 'Държава' in cells[0]:
                continue
                
            # Country - Nationality Table: Държава | Türkçe | Erkek | Kadın | Çoğul
            if len(cells) == 5 and is_cyrillic(cells[0]) and not is_cyrillic(cells[1]) and is_cyrillic(cells[2]):
                country_bg = cells[0]
                country_tr = cells[1]
                man = cells[2]
                woman = cells[3]
                plural = cells[4]
                
                # Questions for country
                add_q(f"{country_tr} ülkesinin Bulgarcası:", country_bg, "Ülke adını yazınız")
                add_q(f"{country_bg} ülkesinin Türkçesi:", country_tr, "Ülke adını yazınız")
                
                # Questions for nationalities
                add_q(f"{country_tr} vatandaşı (erkek):", man, f"{country_tr} erkek milliyet", "Erkekler için tekil milliyet formatı")
                add_q(f"{country_tr} vatandaşı (kadın):", woman, f"{country_tr} kadın milliyet", "Kadınlar için tekil milliyet formatı")
                add_q(f"{country_tr} vatandaşları (çoğul):", plural, f"{country_tr} milliyet çoğul", "Çoğul milliyet formatı")
                
            # Profession / Grammar Table: Български | Турски
            elif len(cells) >= 2 and len(cells) <= 3:
                bg_col = cells[0]
                tr_col = cells[1]
                
                if is_cyrillic(bg_col) and not is_cyrillic(bg_col) and is_cyrillic(tr_col):
                    # sometimes columns might be swapped?
                    bg_col, tr_col = tr_col, bg_col
                    
                if is_cyrillic(bg_col) and tr_col:
                    exp = None
                    if len(cells) == 3 and not is_cyrillic(cells[2]):
                        exp = cells[2]
                    
                    add_q(f"Çeviriniz (Türkçesi: {tr_col})", bg_col, "Bulgarca karşılığını yazınız", exp)
                    add_q(f"Çeviriniz (Bulgarcası: {bg_col})", tr_col, "Türkçe karşılığını yazınız", exp)

parse_docx('D:\\bulgarca_sınav_modulu\\docs\\balgoc___Bulgarca_A1_Ders_3_Turkce_Aciklamali_Not.docx')

print(f"Generated {len(questions)} questions")

out_file = 'src/data/modules/balgoc___Bulgarca_A1_Ders_3.json'
with open(out_file, 'w', encoding='utf-8') as f:
    json.dump({
        "title": "Ders 3: Откъде сте? (Nerelisiniz?)",
        "questions": questions
    }, f, ensure_ascii=False, indent=2)

print(f"Generated {len(questions)} questions in {out_file}")
