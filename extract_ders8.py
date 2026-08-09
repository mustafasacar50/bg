import json
import hashlib
import docx

questions = []

def add_q(sentence, answer, hint, explanation=None, lesson="Ders 8", type_override=None):
    # Unique ID
    raw = sentence + answer + hint + str(type_override)
    qid = "q_auto_" + hashlib.md5(raw.encode('utf-8')).hexdigest()[:8]
    q_type = type_override if type_override else ("fill_in_the_blank" if "_____" in sentence else "translation")
    q = {
        "id": qid,
        "type": q_type,
        "sentence": sentence,
        "answer": answer,
        "hint": hint,
        "lesson": lesson
    }
    if explanation:
        q["explanation"] = explanation
    
    # Avoid exact duplicates
    if not any(x['sentence'] == q['sentence'] and x['answer'] == q['answer'] and x['type'] == q['type'] for x in questions):
        questions.append(q)

def parse_docx(file_path):
    doc = docx.Document(file_path)
    
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    i = 0
    while i < len(paragraphs):
        p = paragraphs[i]
        
        # In Ders 8, Bulgarian text often starts with "—" or just cyrillic, followed by "Türkçesi:"
        # Let's check if the next line starts with "Türkçesi:" or "Türkçe:"
        if i + 1 < len(paragraphs) and paragraphs[i+1].startswith('Türkçe'):
            bg_text = p
            # clean up leading dash
            if bg_text.startswith('—'):
                bg_text = bg_text[1:].strip()
                
            tr_text = paragraphs[i+1].replace('Türkçesi:', '').replace('Türkçe:', '').strip()
            i += 1
            
            exp = None
            if i + 1 < len(paragraphs) and paragraphs[i+1].startswith('Not:'):
                exp = paragraphs[i+1][4:].strip()
                i += 1
                
            # BG to TR
            add_q(
                sentence=bg_text,
                answer=tr_text,
                hint="Türkçe karşılığını yazınız",
                explanation=exp
            )
            # TR to BG
            add_q(
                sentence=tr_text,
                answer=bg_text,
                hint="Bulgarca karşılığını yazınız",
                explanation=exp
            )
            if len(bg_text.split()) >= 2:
                add_q(
                    sentence=tr_text,
                    answer=bg_text,
                    hint="Doğru sıralamayı bul",
                    explanation=exp,
                    type_override="scramble"
                )
        
        i += 1

    # Also check tables for vocab
    for table in doc.tables:
        for r_idx in range(len(table.rows)):
            cells = table.rows[r_idx].cells
            if len(cells) < 2: continue
            
            bg_col = cells[0].text.strip().replace('\n', ' ')
            tr_col = cells[1].text.strip().replace('\n', ' ')
            
            if not bg_col or not tr_col or bg_col.lower() == 'bulgarca' or tr_col.lower() == 'türkçe':
                continue
                
            exp = None
            if len(bg_col) > 1 and len(tr_col) > 1:
                add_q(tr_col, bg_col, "Bulgarca karşılığını yazınız", exp)
                add_q(bg_col, tr_col, "Türkçe karşılığını yazınız", exp)
                if len(bg_col.split()) >= 2:
                    add_q(tr_col, bg_col, "Doğru sıralamayı bul", exp, type_override="scramble")

parse_docx(r'D:\bulgarca_sınav_modulu\docs\balgoc___Bulgarca_A1_Ders_8_Tam_Ceviri_Aciklamali.docx')

output_path = 'src/data/modules/balgoc___Bulgarca_A1_Ders_8.json'
with open(output_path, 'w', encoding='utf-8') as f:
    json.dump({
        "id": "balgoc___Bulgarca_A1_Ders_8",
        "title": "Ders 8: Nerede Yaşıyorsun?",
        "description": "Diyaloglar, mekan kelimeleri, sayılar ve dilbilgisi notları.",
        "questions": questions
    }, f, ensure_ascii=False, indent=2)

print(f"Generated {len(questions)} questions in {output_path}")
