import docx
import json
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc_path = r'D:\bulgarca_sınav_modulu\docs\balgoc___Bulgarca_A1_Ders_4_Turkce_Aciklamali_Not (1).docx'
doc = docx.Document(doc_path)

# Paragrafları yaz
paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
with open('ders4_paragraphs.txt', 'w', encoding='utf-8') as f:
    f.write('\n'.join(paragraphs))

# Tabloları çıkar
all_tables = []
for i, table in enumerate(doc.tables):
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        unique_cells = list(dict.fromkeys(cells))  # Duplicate hücreleri kaldır
        if any(c for c in unique_cells):
            rows.append(unique_cells)
    if rows:
        all_tables.append({"table_index": i, "rows": rows})

with open('ders4_tables_full.json', 'w', encoding='utf-8') as f:
    json.dump(all_tables, f, ensure_ascii=False, indent=2)

print(f"Paragraphs: {len(paragraphs)}")
print(f"Tables: {len(all_tables)}")
for t in all_tables:
    print(f"  Table {t['table_index']}: {len(t['rows'])} rows | Header: {t['rows'][0]}")
